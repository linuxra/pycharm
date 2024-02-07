from docx import Document
import nbformat as nbf
import os


def save_image(image, image_count):
    image_name = f"image{image_count}.jpg"
    with open(image_name, 'wb') as f:
        f.write(image.blob)
    return image_name


def docx_to_markdown(doc_path):
    doc = Document(doc_path)
    markdown = []
    image_count = 0

    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            level = len(para.style.name.split('Heading ')[1])
            markdown.append('#' * level + ' ' + para.text)
        else:
            markdown.append(para.text)

    for table in doc.tables:
        markdown_table = ["| " + " | ".join(cell.text for cell in row.cells) + " |" for row in table.rows]
        header_separator = ["|:---" * len(table.rows[0].cells) + "|"]
        markdown.extend(markdown_table[:1] + header_separator + markdown_table[1:])

    for shape in [s for p in doc.paragraphs for s in p.runs if s.element.find('./w:drawing/') is not None]:
        image = shape.element.find('./w:drawing/').getchildren()[0].getchildren()[0].getchildren()[0].getchildren()[
            0].getchildren()[1].embed
        image_file = doc.part.related_parts[image].blob
        image_name = save_image(image_file, image_count)
        markdown.append(f"![Image](./{image_name})")
        image_count += 1

    return markdown


def create_notebook(markdown_cells, output_path='output_notebook.ipynb'):
    nb = nbf.v4.new_notebook()
    for cell_content in markdown_cells:
        cell = nbf.v4.new_markdown_cell(cell_content)
        nb['cells'].append(cell)

    with open(output_path, 'w') as f:
        nbf.write(nb, f)


# Usage
markdown_cells = docx_to_markdown('path/to/your/document.docx')
create_notebook(markdown_cells)
from docx import Document
from io import BytesIO
from PIL import Image


def extract_images(word_doc_path):
  """
  Extracts images from a Word document and saves them to a specified folder.

  Args:
    word_doc_path: Path to the Word document file.

  Returns:
    A list of extracted image filenames.
  """
  doc = Document(word_doc_path)
  extracted_images = []

  for rel in doc.part.rels.values():
    if rel.reltype == docx.opc.RELATIONSHIP_TYPE.IMAGE:
      image_filename = rel.target
      image_data = doc.part.related_parts[image_filename].blob

      # Save the image to a file
      img = Image.open(BytesIO(image_data))
      img_filename = f"extracted_image_{len(extracted_images)}.png"
      img.save(img_filename)

      extracted_images.append(img_filename)

  return extracted_images


if __name__ == "__main__":
  # Example usage
  word_doc_path = "your_word_document.docx"
  extracted_filenames = extract_images(word_doc_path)

  print(f"Extracted {len(extracted_filenames)} images:")
  for filename in extracted_filenames:
    print(filename)
